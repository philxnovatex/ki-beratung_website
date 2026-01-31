# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import os

# Pfade
script_dir = r"c:\Users\AIcon\OneDrive\Desktop\KI_Beratung\01_Website\KI-Beratung_Website\documents\templates"
logo_path = r"c:\Users\AIcon\OneDrive\Desktop\KI_Beratung\01_Website\KI-Beratung_Website\public\assets\images\neuratex-logo.jpg"
output_path = os.path.join(script_dir, "OnePager_Kraemer24.docx")

# Farben
GOLD_HEX = "FFC947"
BLUE_HEX = "185ADB"
DARK_BLUE_HEX = "0A1931"
GREEN_HEX = "E8F5E9"
GREEN_TEXT_HEX = "16A34A"
YELLOW_HEX = "FFF8E1"
ORANGE_TEXT_HEX = "B45309"
LIGHT_BLUE_HEX = "F0F4FF"
GRAY_HEX = "F8F9FA"

def set_cell_shading(cell, color_hex):
    """Setzt Hintergrundfarbe einer Zelle"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, border_color="E0E0E0", left_color=None, left_size=24):
    """Setzt Rahmen einer Zelle mit optionalem dicken linken Rand"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ['top', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    
    # Linker Rand dicker und farbig
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(left_size))
    left.set(qn('w:color'), left_color if left_color else GOLD_HEX)
    tcBorders.append(left)
    
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Setzt Innenabstand einer Zelle"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(margin_val))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def add_styled_text(paragraph, text, size=10, color_hex=DARK_BLUE_HEX, bold=False):
    """Fuegt formatierten Text hinzu"""
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16))
    run.bold = bold
    return run

# Dokument erstellen
doc = Document()

# Seitenraender
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ============================================
# HEADER TABELLE
# ============================================
header_table = doc.add_table(rows=1, cols=2)
header_table.autofit = True
header_table.columns[0].width = Inches(3)
header_table.columns[1].width = Inches(3.5)

# Logo
cell_logo = header_table.cell(0, 0)
if os.path.exists(logo_path):
    p = cell_logo.paragraphs[0]
    run = p.add_run()
    run.add_picture(logo_path, height=Inches(0.55))

# Rechts: One Pager
cell_right = header_table.cell(0, 1)
p = cell_right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_styled_text(p, "ONE PAGER", 10, GOLD_HEX, bold=True)
p.add_run("\n")
add_styled_text(p, "Januar 2025", 9, "666666")

# Gold-Linie unter Header
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
# Horizontale Linie als Tabelle
line_table = doc.add_table(rows=1, cols=1)
line_table.cell(0, 0).width = Inches(6.5)
set_cell_shading(line_table.cell(0, 0), GOLD_HEX)
line_cell = line_table.cell(0, 0)
line_cell.paragraphs[0].clear()
# Hoehe der Linie
tr = line_table.rows[0]._tr
trPr = tr.get_or_add_trPr()
trHeight = OxmlElement('w:trHeight')
trHeight.set(qn('w:val'), '60')
trHeight.set(qn('w:hRule'), 'exact')
trPr.append(trHeight)

doc.add_paragraph()

# ============================================
# TITEL
# ============================================
p = doc.add_paragraph()
add_styled_text(p, "Internes Wissens- und Dokumentensystem", 22, DARK_BLUE_HEX, bold=True)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
add_styled_text(p, "Kraemer24", 14, BLUE_HEX, bold=False)
p.paragraph_format.space_after = Pt(16)

# ============================================
# AUSGANGSLAGE
# ============================================
p = doc.add_paragraph()
add_styled_text(p, "Ausgangslage", 13, BLUE_HEX, bold=True)
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
add_styled_text(p, "Im Ersatzteilgeschaeft von ", 10, DARK_BLUE_HEX)
add_styled_text(p, "Kraemer24", 10, DARK_BLUE_HEX, bold=True)
add_styled_text(p, " liegen relevante Informationen verteilt in unterschiedlichen Systemen und Formaten:", 10, DARK_BLUE_HEX)
p.paragraph_format.space_after = Pt(6)

for item in [
    "PDFs, gescannte Kataloge, Excel-Listen, CDs und DVDs",
    "Herstellerunterlagen mit sehr unterschiedlicher Struktur und Qualitaet",
    "Erfahrungswissen einzelner Mitarbeiter zu Fabrikaten, Baujahren und Sonderfaellen"
]:
    p = doc.add_paragraph()
    add_styled_text(p, "   " + item, 10, DARK_BLUE_HEX)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)

p = doc.add_paragraph()
add_styled_text(p, "Die Identifikation von Maschinen und passenden Ersatzteilen erfordert haeufig manuelle Recherche und persoenliche Erfahrung. Dies fuehrt zu erhoehtem Zeitaufwand im Kundenservice und macht die Trefferquote stark abhaengig von einzelnen Schluesselpersonen.", 10, DARK_BLUE_HEX)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
add_styled_text(p, "Bestehende Werkzeuge wie Microsoft Teams und SharePoint ermoeglichen zwar Austausch, ersetzen jedoch kein strukturiertes Wissenssystem und keine standardisierte Fallnachverfolgung.", 10, DARK_BLUE_HEX)
p.paragraph_format.space_after = Pt(14)

# ============================================
# ZIELBILD BOX
# ============================================
p = doc.add_paragraph()
add_styled_text(p, "Zielbild", 13, BLUE_HEX, bold=True)
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
add_styled_text(p, "Ziel ist kein weiteres Ablagesystem, sondern ein internes Werkzeug, das:", 10, DARK_BLUE_HEX)
p.paragraph_format.space_after = Pt(8)

# Zielbild als 2x2 Tabelle
ziel_table = doc.add_table(rows=2, cols=2)
ziel_table.autofit = True

ziel_items = [
    ("Suche", "Informationen schnell auffindbar macht"),
    ("Wissen", "Erfahrungswissen sichert"),
    ("Start", "Neue Mitarbeiter schneller handlungsfaehig macht"),
    ("Speed", "Den Ersatzteilservice spuerbar entlastet")
]

for i, (icon, text) in enumerate(ziel_items):
    row = i // 2
    col = i % 2
    cell = ziel_table.cell(row, col)
    set_cell_shading(cell, LIGHT_BLUE_HEX)
    set_cell_border(cell, "D0D8E8", BLUE_HEX, 16)
    set_cell_margins(cell, 80, 80, 120, 120)
    p = cell.paragraphs[0]
    add_styled_text(p, "[" + icon + "]  ", 9, BLUE_HEX, bold=True)
    add_styled_text(p, text, 10, DARK_BLUE_HEX)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
add_styled_text(p, "Das System soll intern genutzt werden, den Arbeitsalltag vereinfachen und schrittweise eingefuehrt werden, um Akzeptanz im Team zu gewaehrleisten.", 10, DARK_BLUE_HEX)
p.paragraph_format.space_after = Pt(14)

# ============================================
# DREI HEBEL
# ============================================
p = doc.add_paragraph()
add_styled_text(p, "Drei Hebel zur Verbesserung", 13, BLUE_HEX, bold=True)
p.paragraph_format.space_after = Pt(10)

def add_hebel_box(doc, number, title, description, nutzen_items, wichtig=False):
    """Erstellt eine Hebel-Box als Tabelle"""
    box = doc.add_table(rows=1, cols=1)
    box.autofit = True
    cell = box.cell(0, 0)
    set_cell_shading(cell, GRAY_HEX)
    set_cell_border(cell, "E0E0E0", GOLD_HEX, 24)
    set_cell_margins(cell, 120, 120, 150, 150)
    
    # Titel
    p = cell.paragraphs[0]
    add_styled_text(p, f"[{number}] ", 11, BLUE_HEX, bold=True)
    add_styled_text(p, title, 12, DARK_BLUE_HEX, bold=True)
    
    # Beschreibung
    p = cell.add_paragraph()
    add_styled_text(p, description, 10, DARK_BLUE_HEX)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    
    # Nutzen-Box innerhalb
    nutzen_table = cell.add_table(rows=1, cols=1)
    nutzen_cell = nutzen_table.cell(0, 0)
    bg_color = YELLOW_HEX if wichtig else GREEN_HEX
    set_cell_shading(nutzen_cell, bg_color)
    set_cell_margins(nutzen_cell, 80, 80, 100, 100)
    
    # Nutzen-Titel
    p = nutzen_cell.paragraphs[0]
    label = "WICHTIG" if wichtig else "NUTZEN"
    label_color = ORANGE_TEXT_HEX if wichtig else GREEN_TEXT_HEX
    add_styled_text(p, label, 8, label_color, bold=True)
    
    # Nutzen-Items
    for item in nutzen_items:
        p = nutzen_cell.add_paragraph()
        add_styled_text(p, "- " + item, 9, DARK_BLUE_HEX)
        p.paragraph_format.space_after = Pt(2)
    
    # Abstand nach Box
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

add_hebel_box(doc, "1", "Altunterlagen nutzbar machen",
    "Historische Dokumente wie Kataloge, Scans und Tabellen werden so aufbereitet, dass Inhalte zuverlaessig durchsuchbar und strukturiert nutzbar sind.",
    ["Schnellere Recherche bei Altmaschinen und Sonderfaellen",
     "Weniger Abhaengigkeit von Papier und Einzelwissen",
     "Basis fuer weitere Automatisierung"])

add_hebel_box(doc, "2", "Erfahrungswissen im Team sichern",
    "Reale Loesungsfaelle aus dem Alltag werden in kompakter Form dokumentiert und fuer alle Mitarbeiter zugaenglich gemacht.",
    ["Gleichbleibende Antwortqualitaet",
     "Schnellere Einarbeitung neuer Kollegen",
     "Wissen bleibt im Unternehmen, auch bei Personalwechsel"])

# ============================================
# SEITE 2
# ============================================
doc.add_page_break()

# Header Seite 2
header2 = doc.add_table(rows=1, cols=2)
header2.columns[0].width = Inches(3)
header2.columns[1].width = Inches(3.5)

if os.path.exists(logo_path):
    p = header2.cell(0, 0).paragraphs[0]
    run = p.add_run()
    run.add_picture(logo_path, height=Inches(0.55))

p = header2.cell(0, 1).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_styled_text(p, "ONE PAGER", 10, GOLD_HEX, bold=True)
p.add_run("\n")
add_styled_text(p, "Januar 2025", 9, "666666")

# Gold-Linie
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
line2 = doc.add_table(rows=1, cols=1)
set_cell_shading(line2.cell(0, 0), GOLD_HEX)
line2.cell(0, 0).paragraphs[0].clear()
tr2 = line2.rows[0]._tr
trPr2 = tr2.get_or_add_trPr()
trH2 = OxmlElement('w:trHeight')
trH2.set(qn('w:val'), '60')
trH2.set(qn('w:hRule'), 'exact')
trPr2.append(trH2)

doc.add_paragraph()

# Hebel 3
add_hebel_box(doc, "3", "Schneller Zugriff ueber internes Assistenzsystem",
    "Ein internes Assistenzsystem ermoeglicht es, Fragen auf Basis der vorhandenen Dokumente und Wissensfaelle zu beantworten.",
    ["Antworten basieren ausschliesslich auf internen Quellen",
     "Jede Antwort ist nachvollziehbar und referenziert",
     "Kein Einsatz im Endkundensupport"],
    wichtig=True)

# Extra Nutzen fuer Hebel 3
nutzen_extra = doc.add_table(rows=1, cols=1)
nutzen_extra_cell = nutzen_extra.cell(0, 0)
set_cell_shading(nutzen_extra_cell, GREEN_HEX)
set_cell_margins(nutzen_extra_cell, 80, 80, 100, 100)
p = nutzen_extra_cell.paragraphs[0]
add_styled_text(p, "NUTZEN", 8, GREEN_TEXT_HEX, bold=True)
for item in ["Deutlich reduzierte Suchzeiten", "Weniger interne Rueckfragen", "Entlastung erfahrener Mitarbeiter"]:
    p = nutzen_extra_cell.add_paragraph()
    add_styled_text(p, "- " + item, 9, DARK_BLUE_HEX)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# ============================================
# VORGEHENSMODELL
# ============================================
p = doc.add_paragraph()
add_styled_text(p, "Vorgehensmodell", 13, BLUE_HEX, bold=True)
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
add_styled_text(p, "Kein Grossprojekt. Kein Systemwechsel.", 11, DARK_BLUE_HEX, bold=True)
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
add_styled_text(p, "Vorgeschlagen wird ein klar abgegrenzter Machbarkeitstest:", 10, DARK_BLUE_HEX)
p.paragraph_format.space_after = Pt(8)

# Schritte als Box
steps_box = doc.add_table(rows=1, cols=1)
steps_cell = steps_box.cell(0, 0)
set_cell_shading(steps_cell, LIGHT_BLUE_HEX)
set_cell_border(steps_cell, "D0D8E8", BLUE_HEX, 16)
set_cell_margins(steps_cell, 100, 100, 120, 120)

p = steps_cell.paragraphs[0]
add_styled_text(p, "Schritt fuer Schritt", 11, BLUE_HEX, bold=True)

for i, step in enumerate([
    "Nutzung realer, komplexer Unterlagen",
    "Test an typischen Spezialfaellen aus dem Alltag",
    "Bewertung von Nutzen, Qualitaet und Akzeptanz"
], 1):
    p = steps_cell.add_paragraph()
    add_styled_text(p, f"{i}. {step}", 10, DARK_BLUE_HEX)
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
add_styled_text(p, "Ergebnis ist eine belastbare Entscheidungsgrundlage, ob und in welchem Umfang eine Weiterentwicklung sinnvoll ist.", 10, DARK_BLUE_HEX)
p.paragraph_format.space_after = Pt(12)

# ============================================
# ENTSCHEIDUNGSFRAGEN
# ============================================
p = doc.add_paragraph()
add_styled_text(p, "Entscheidungsfragen fuer die Fachbereiche", 13, BLUE_HEX, bold=True)
p.paragraph_format.space_after = Pt(8)

questions = [
    "Wo entstehen heute die groessten Zeitverluste bei der Ersatzteilsuche?",
    "Welche Unterlagen oder Fabrikate sind besonders problematisch?",
    "Wie hoch ist die Abhaengigkeit von einzelnen Mitarbeitern?",
    "Welche Loesung wuerde den Alltag messbar erleichtern?"
]

q_table = doc.add_table(rows=2, cols=2)
for i, q in enumerate(questions):
    row = i // 2
    col = i % 2
    cell = q_table.cell(row, col)
    set_cell_shading(cell, GRAY_HEX)
    set_cell_border(cell, "E0E0E0", BLUE_HEX, 12)
    set_cell_margins(cell, 80, 80, 100, 100)
    p = cell.paragraphs[0]
    add_styled_text(p, "[?] ", 10, BLUE_HEX, bold=True)
    add_styled_text(p, q, 9, DARK_BLUE_HEX)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# ============================================
# NAECHSTER SCHRITT BOX
# ============================================
cta_box = doc.add_table(rows=1, cols=1)
cta_cell = cta_box.cell(0, 0)
set_cell_shading(cta_cell, LIGHT_BLUE_HEX)
set_cell_border(cta_cell, GOLD_HEX, GOLD_HEX, 24)
set_cell_margins(cta_cell, 120, 120, 150, 150)

p = cta_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_styled_text(p, "Naechster Schritt", 13, BLUE_HEX, bold=True)

p = cta_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_styled_text(p, "Sammlung von 1-2 konkreten Anwendungsfaellen und Beispielunterlagen aus Produktmanagement und Vertrieb, um die technische und organisatorische Machbarkeit gezielt zu pruefen.", 10, DARK_BLUE_HEX)

doc.add_paragraph()

# ============================================
# FOOTER
# ============================================
p = doc.add_paragraph()
add_styled_text(p, "neuratex.de | philippkoch@neuratex.de", 9, "666666")

# Speichern
doc.save(output_path)
print(f"Word-Dokument erstellt: {output_path}")
