# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Pfade
script_dir = r"c:\Users\AIcon\OneDrive\Desktop\KI_Beratung\01_Website\KI-Beratung_Website\documents\templates"
logo_path = r"c:\Users\AIcon\OneDrive\Desktop\KI_Beratung\01_Website\KI-Beratung_Website\public\assets\images\neuratex-logo.jpg"
output_path = os.path.join(script_dir, "OnePager_Kraemer24.docx")

# Farben
GOLD = "FFC947"
BLUE = "185ADB"
DARK = "0A1931"
GREEN_BG = "E8F5E9"
GREEN_TEXT = "16A34A"
YELLOW_BG = "FFF8E1"
ORANGE_TEXT = "B45309"
LIGHT_BLUE = "F0F4FF"
GRAY = "F8F9FA"

def rgb(hex_color):
    return RGBColor(int(hex_color[:2], 16), int(hex_color[2:4], 16), int(hex_color[4:], 16))

def set_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_borders(cell, left_color=None, left_size=24):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'E0E0E0')
        borders.append(b)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(left_size))
    left.set(qn('w:color'), left_color or GOLD)
    borders.append(left)
    tcPr.append(borders)

def set_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{name}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)

def text(p, txt, size=10, color=DARK, bold=False):
    run = p.add_run(txt)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    return run

def keep_together(paragraph):
    """Verhindert Seitenumbruch innerhalb des Absatzes"""
    pPr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    keep.set(qn('w:val'), '1')
    pPr.append(keep)
    keep2 = OxmlElement('w:keepLines')
    keep2.set(qn('w:val'), '1')
    pPr.append(keep2)

# Dokument
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(2)
    s.right_margin = Cm(2)

# === HEADER ===
h = doc.add_table(rows=1, cols=2)
if os.path.exists(logo_path):
    h.cell(0,0).paragraphs[0].add_run().add_picture(logo_path, height=Inches(0.5))
p = h.cell(0,1).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
text(p, "ONE PAGER", 10, GOLD, True)
p.add_run("\n")
text(p, "Januar 2025", 9, "666666")

# Gold-Linie
doc.add_paragraph()
line = doc.add_table(rows=1, cols=1)
set_shading(line.cell(0,0), GOLD)
line.cell(0,0).paragraphs[0].clear()
tr = line.rows[0]._tr
trPr = tr.get_or_add_trPr()
trH = OxmlElement('w:trHeight')
trH.set(qn('w:val'), '50')
trH.set(qn('w:hRule'), 'exact')
trPr.append(trH)

doc.add_paragraph()

# === TITEL ===
p = doc.add_paragraph()
text(p, "Internes Wissens- und Dokumentensystem", 22, DARK, True)
keep_together(p)

p = doc.add_paragraph()
text(p, "Kraemer24", 14, BLUE)
p.paragraph_format.space_after = Pt(14)

# === AUSGANGSLAGE ===
p = doc.add_paragraph()
text(p, "Ausgangslage", 13, BLUE, True)
p.paragraph_format.space_after = Pt(6)
keep_together(p)

p = doc.add_paragraph()
text(p, "Im Ersatzteilgeschäft von ", 10, DARK)
text(p, "Kraemer24", 10, DARK, True)
text(p, " liegen relevante Informationen verteilt in unterschiedlichen Systemen und Formaten:", 10, DARK)
p.paragraph_format.space_after = Pt(6)

for item in [
    "PDFs, gescannte Kataloge, Excel-Listen, CDs und DVDs",
    "Herstellerunterlagen mit sehr unterschiedlicher Struktur und Qualität",
    "Erfahrungswissen einzelner Mitarbeiter zu Fabrikaten, Baujahren und Sonderfällen"
]:
    p = doc.add_paragraph()
    text(p, "• " + item, 10, DARK)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
text(p, "Die Identifikation von Maschinen und passenden Ersatzteilen erfordert häufig manuelle Recherche und persönliche Erfahrung. Dies führt zu erhöhtem Zeitaufwand im Kundenservice und macht die Trefferquote stark abhängig von einzelnen Schlüsselpersonen.", 10, DARK)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
text(p, "Bestehende Werkzeuge wie Microsoft Teams und SharePoint ermöglichen zwar Austausch, ersetzen jedoch kein strukturiertes Wissenssystem und keine standardisierte Fallnachverfolgung.", 10, DARK)
p.paragraph_format.space_after = Pt(12)

# === ZIELBILD ===
p = doc.add_paragraph()
text(p, "Zielbild", 13, BLUE, True)
p.paragraph_format.space_after = Pt(6)
keep_together(p)

p = doc.add_paragraph()
text(p, "Ziel ist kein weiteres Ablagesystem, sondern ein internes Werkzeug, das:", 10, DARK)
p.paragraph_format.space_after = Pt(8)

# 2x2 Grid
zt = doc.add_table(rows=2, cols=2)
ziel = [
    ("🔍", "Informationen schnell auffindbar macht"),
    ("🧠", "Erfahrungswissen sichert"),
    ("🚀", "Neue Mitarbeiter schneller handlungsfähig macht"),
    ("⚡", "Den Ersatzteilservice spürbar entlastet")
]
for i, (icon, txt) in enumerate(ziel):
    c = zt.cell(i//2, i%2)
    set_shading(c, LIGHT_BLUE)
    set_borders(c, BLUE, 16)
    set_margins(c, 80, 80, 120, 120)
    p = c.paragraphs[0]
    text(p, icon + "  ", 11, DARK)
    text(p, txt, 10, DARK)

doc.add_paragraph()
p = doc.add_paragraph()
text(p, "Das System soll intern genutzt werden, den Arbeitsalltag vereinfachen und schrittweise eingeführt werden, um Akzeptanz im Team zu gewährleisten.", 10, DARK)
p.paragraph_format.space_after = Pt(12)

# === DREI HEBEL ===
p = doc.add_paragraph()
text(p, "Drei Hebel zur Verbesserung", 13, BLUE, True)
p.paragraph_format.space_after = Pt(8)
keep_together(p)

def hebel(doc, num, title, desc, items, wichtig=False):
    box = doc.add_table(rows=1, cols=1)
    c = box.cell(0,0)
    set_shading(c, GRAY)
    set_borders(c, GOLD, 24)
    set_margins(c, 100, 100, 140, 140)
    
    p = c.paragraphs[0]
    text(p, f"[{num}] ", 11, BLUE, True)
    text(p, title, 11, DARK, True)
    
    p = c.add_paragraph()
    text(p, desc, 10, DARK)
    p.paragraph_format.space_after = Pt(8)
    
    # Nutzen-Box
    nt = c.add_table(rows=1, cols=1)
    nc = nt.cell(0,0)
    set_shading(nc, YELLOW_BG if wichtig else GREEN_BG)
    set_margins(nc, 60, 60, 80, 80)
    
    p = nc.paragraphs[0]
    label = "WICHTIG" if wichtig else "NUTZEN"
    color = ORANGE_TEXT if wichtig else GREEN_TEXT
    text(p, label, 8, color, True)
    
    for item in items:
        p = nc.add_paragraph()
        text(p, "• " + item, 9, DARK)
        p.paragraph_format.space_after = Pt(1)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

hebel(doc, "1", "Altunterlagen nutzbar machen",
    "Historische Dokumente wie Kataloge, Scans und Tabellen werden so aufbereitet, dass Inhalte zuverlässig durchsuchbar und strukturiert nutzbar sind.",
    ["Schnellere Recherche bei Altmaschinen und Sonderfällen",
     "Weniger Abhängigkeit von Papier und Einzelwissen",
     "Basis für weitere Automatisierung"])

hebel(doc, "2", "Erfahrungswissen im Team sichern",
    "Reale Lösungsfälle aus dem Alltag werden in kompakter Form dokumentiert und für alle Mitarbeiter zugänglich gemacht.",
    ["Gleichbleibende Antwortqualität",
     "Schnellere Einarbeitung neuer Kollegen",
     "Wissen bleibt im Unternehmen, auch bei Personalwechsel"])

# === SEITE 2 ===
doc.add_page_break()

# Header S2
h2 = doc.add_table(rows=1, cols=2)
if os.path.exists(logo_path):
    h2.cell(0,0).paragraphs[0].add_run().add_picture(logo_path, height=Inches(0.5))
p = h2.cell(0,1).paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
text(p, "ONE PAGER", 10, GOLD, True)
p.add_run("\n")
text(p, "Januar 2025", 9, "666666")

doc.add_paragraph()
line2 = doc.add_table(rows=1, cols=1)
set_shading(line2.cell(0,0), GOLD)
line2.cell(0,0).paragraphs[0].clear()
tr2 = line2.rows[0]._tr
trPr2 = tr2.get_or_add_trPr()
trH2 = OxmlElement('w:trHeight')
trH2.set(qn('w:val'), '50')
trH2.set(qn('w:hRule'), 'exact')
trPr2.append(trH2)
doc.add_paragraph()

# Hebel 3
hebel(doc, "3", "Schneller Zugriff über internes Assistenzsystem",
    "Ein internes Assistenzsystem ermöglicht es, Fragen auf Basis der vorhandenen Dokumente und Wissensfälle zu beantworten.",
    ["Antworten basieren ausschließlich auf internen Quellen",
     "Jede Antwort ist nachvollziehbar und referenziert",
     "Kein Einsatz im Endkundensupport"],
    wichtig=True)

# Extra Nutzen
nt = doc.add_table(rows=1, cols=1)
nc = nt.cell(0,0)
set_shading(nc, GREEN_BG)
set_margins(nc, 60, 60, 80, 80)
p = nc.paragraphs[0]
text(p, "NUTZEN", 8, GREEN_TEXT, True)
for item in ["Deutlich reduzierte Suchzeiten", "Weniger interne Rückfragen", "Entlastung erfahrener Mitarbeiter"]:
    p = nc.add_paragraph()
    text(p, "• " + item, 9, DARK)
doc.add_paragraph().paragraph_format.space_after = Pt(10)

# === VORGEHENSMODELL ===
p = doc.add_paragraph()
text(p, "Vorgehensmodell", 13, BLUE, True)
p.paragraph_format.space_after = Pt(6)
keep_together(p)

p = doc.add_paragraph()
text(p, "Kein Großprojekt. Kein Systemwechsel.", 11, DARK, True)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
text(p, "Vorgeschlagen wird ein klar abgegrenzter Machbarkeitstest:", 10, DARK)
p.paragraph_format.space_after = Pt(6)

# Schritte Box
sb = doc.add_table(rows=1, cols=1)
sc = sb.cell(0,0)
set_shading(sc, LIGHT_BLUE)
set_borders(sc, BLUE, 16)
set_margins(sc, 80, 80, 100, 100)
p = sc.paragraphs[0]
text(p, "Schritt für Schritt", 11, BLUE, True)
for i, step in enumerate(["Nutzung realer, komplexer Unterlagen", "Test an typischen Spezialfällen aus dem Alltag", "Bewertung von Nutzen, Qualität und Akzeptanz"], 1):
    p = sc.add_paragraph()
    text(p, f"{i}. {step}", 10, DARK)

doc.add_paragraph()
p = doc.add_paragraph()
text(p, "Ergebnis ist eine belastbare Entscheidungsgrundlage, ob und in welchem Umfang eine Weiterentwicklung sinnvoll ist.", 10, DARK)
p.paragraph_format.space_after = Pt(10)

# === ENTSCHEIDUNGSFRAGEN ===
p = doc.add_paragraph()
text(p, "Entscheidungsfragen für die Fachbereiche", 13, BLUE, True)
p.paragraph_format.space_after = Pt(6)
keep_together(p)

qt = doc.add_table(rows=2, cols=2)
qs = [
    "Wo entstehen heute die größten Zeitverluste bei der Ersatzteilsuche?",
    "Welche Unterlagen oder Fabrikate sind besonders problematisch?",
    "Wie hoch ist die Abhängigkeit von einzelnen Mitarbeitern?",
    "Welche Lösung würde den Alltag messbar erleichtern?"
]
for i, q in enumerate(qs):
    c = qt.cell(i//2, i%2)
    set_shading(c, GRAY)
    set_borders(c, BLUE, 12)
    set_margins(c, 60, 60, 80, 80)
    p = c.paragraphs[0]
    text(p, "? ", 10, BLUE, True)
    text(p, q, 9, DARK)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# === CTA ===
cta = doc.add_table(rows=1, cols=1)
cc = cta.cell(0,0)
set_shading(cc, LIGHT_BLUE)
set_borders(cc, GOLD, 24)
set_margins(cc, 100, 100, 120, 120)
p = cc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
text(p, "🚀 Nächster Schritt", 12, BLUE, True)
p = cc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
text(p, "Sammlung von 1–2 konkreten Anwendungsfällen und Beispielunterlagen aus Produktmanagement und Vertrieb, um die technische und organisatorische Machbarkeit gezielt zu prüfen.", 10, DARK)

doc.add_paragraph()

# Footer
p = doc.add_paragraph()
text(p, "neuratex.de | philippkoch@neuratex.de", 9, "666666")

doc.save(output_path)
print("Fertig:", output_path)
